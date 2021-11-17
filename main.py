# 数据分析获取权威排行榜
import re
from tqdm import tqdm
import random
import uuid
import os
import time
from glob import glob
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from urllib.request import urlretrieve
import docx
from docx.oxml.ns import qn
import shutil
import threading
from bs4 import BeautifulSoup, NavigableString, Tag



chrome_options = Options()
chrome_options.add_argument('--headless')
chrome_options.add_argument('--disable-gpu')
driver_path = '../common/webdriver/chromedriver.exe'
#不同类型的操作系统以及浏览器的标识
user_agent_list=[
        'Mozilla/5.0(compatible;MSIE9.0;WindowsNT6.1;Trident/5.0)',
        'Mozilla/4.0(compatible;MSIE8.0;WindowsNT6.0;Trident/4.0)',
        'Mozilla/4.0(compatible;MSIE7.0;WindowsNT6.0)',
        'Opera/9.80(WindowsNT6.1;U;en)Presto/2.8.131Version/11.11',
        'Mozilla/5.0(WindowsNT6.1;rv:2.0.1)Gecko/20100101Firefox/4.0.1',
        'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/21.0.1180.71 Safari/537.1 LBBROWSER',
        'Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; QQDownload 732; .NET4.0C; .NET4.0E)',
        'Mozilla/5.0 (Windows NT 5.1) AppleWebKit/535.11 (KHTML, like Gecko) Chrome/17.0.963.84 Safari/535.11 SE 2.X MetaSr 1.0',
        'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Maxthon/4.4.3.4000 Chrome/30.0.1599.101 Safari/537.36', 
        'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/38.0.2125.122 UBrowser/4.0.3214.0 Safari/537.36',
        'Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/49.0.2623.221 Safari/537.36 SE 2.X MetaSr 1.0'
    ]


def get_web_source_bs(url):
    chrome_options.add_argument('--user-agent=%s' % random.choice(user_agent_list))
    browser = webdriver.Chrome(executable_path=driver_path,chrome_options=chrome_options)
    browser.get(url)
    response = browser.page_source
    bs = BeautifulSoup(response, "html.parser")
    tpg = bs.find_all(["title","p","section","img"])
    return tpg

def get_id():
    return str(uuid.uuid4()).replace("-","")

def get_chinese_str(text):
    res = ""
    for _char in text.strip():
        if '\u4e00' <= _char <= '\u9fa5':
            res += _char
    return res
         
def get_img(imgurl, iid):
    # @ TODO 如何快速爬取单一网址的多个图片
    imgs_folder = "../data/imgs/"
    urlretrieve(imgurl, os.path.join(imgs_folder, str(iid) + '.gif'))
    time.sleep(1)


def get_raw_data(tags):
    global raw_data
    lock = threading.Lock()
    
    for i, t in tags:
        # 获得图片
        data_src = t.find_all(attrs={"data-src":re.compile(r'^http')})
        if data_src:
            lock.acquire()
            raw_data.append((i, ".gif"))
            lock.release()
            get_img(data_src[0]['data-src'], i)
            
        lock.acquire()
        raw_data.append((i, t.text))
        lock.release()
        
        
def get_raw_data_multi_threading(tpg):
    threads = []
    tpg = list(zip(range(len(tpg)), tpg))
    for i in range(0, len(tpg), 5):
        t = threading.Thread(target=get_raw_data, args=(tpg[i: i+5],))
        threads.append(t)
    
    for th in threads:
        th.setDaemon(True)
        th.start()
        

def write_docx(raw_list):
    
    raw_list = sorted(raw_list, key=lambda x: x[0])
    raw_list = [x for x in raw_list if len(x[1].strip())]

    title = get_chinese_str(raw_list[0][1])[:20]
    imgfiles = {os.path.basename(x): x for x in glob("../data/imgs/*.gif")}
    
    doc = docx.Document()
    # 设置正文中文字体
    microsoft_font = u'微软雅黑'  # u 表示后面的字符串以 Unicode 格式进行编码
    area = qn('w:eastAsia')
    doc.styles['Normal'].font.name = microsoft_font
    doc.styles['Normal']._element.rPr.rFonts.set(area, microsoft_font)
    doc.add_heading('', level=1).add_run(raw_list[0][1])
    
    seen = set()
    for line in raw_list[1:]:
        if line[1].startswith("data:image") or not len(line[1]):
            continue
        else:
            if line[1] == ".gif":
                file = str(line[0]) + line[1]
                if file in imgfiles:
                    doc.add_picture(imgfiles[file],width=docx.shared.Cm(15),height=docx.shared.Cm(7))
            else:
                li = line[1]
                if li[: int(len(li)/2)] == li[int(len(li)/2) :]:
                    li = li[: int(len(li)/2)] 
                if li not in seen:
                    doc.add_paragraph(li)
                    seen.add(li)
                    
    doc.save("../data/" + title + get_id() + ".docx")
    
def engin(url):
    shutil.rmtree("../data/imgs/")
    os.mkdir("../data/imgs/")
    bstag = get_web_source_bs(url)
    return bstag
  
if __name__ == "__main__":
  url = "https://mp.weixin.qq.com/s/iYKGeJFVLDBO01f7CpFvUw"
  raw_data = []
  bstag = engin(url)
  get_raw_data_multi_threading(bstag)
  time.sleep(10)
  write_docx(raw_data)
 
